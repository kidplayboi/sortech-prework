"""제출용 마크다운 → docx 변환.

과제가 docx/pdf를 요구한다. pandoc이 없는 환경이라 python-docx로 직접 조립한다.
필요한 것만 지원: 제목(#~###) · 표 · 코드블록 · 인용 · 목록 · 굵게 · 수평선.
화려하게 만들지 않는다 — 읽히는 게 목적이고, 표와 코드블록이 뭉개지지만 않으면 된다.

사용법: python tools/make_docx.py docs/제출물/S1-AI활용기록.md [출력경로]
"""
import pathlib
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

BODY_FONT = "맑은 고딕"
MONO_FONT = "Consolas"
GRAY = RGBColor(0x55, 0x55, 0x55)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")


def add_runs(paragraph, text, base_size=10.5, mono=False):
    """**굵게** 와 `코드` 만 처리한다. 나머지는 그대로."""
    token = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(base_size - 1)
        else:
            run = paragraph.add_run(part)
            if mono:
                run.font.name = MONO_FONT
        if not run.font.name:
            run.font.name = BODY_FONT
        if not run.font.size:
            run.font.size = Pt(base_size)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_hrule(doc):
    """빈 문단에 아래 테두리를 넣어 구분선 하나만 그린다."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    borders.append(bottom)
    pPr.append(borders)


def _is_structural(stripped, lines, idx):
    """문단을 끊어야 하는 줄인가 (제목·표·코드·인용·목록·수평선)"""
    if stripped.startswith(("#", ">", "```")):
        return True
    if re.match(r"^-{3,}$", stripped):
        return True
    if re.match(r"^(\s*)([-*]\s+|\d+\.\s+)", lines[idx]):
        return True
    return (stripped.startswith("|") and idx + 1 < len(lines)
            and re.match(r"^\|[\s:|-]+\|$", lines[idx + 1].strip()) is not None)


def convert(md_path, out_path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 코드블록
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(14)
            para.paragraph_format.space_after = Pt(8)
            run = para.add_run("\n".join(block))
            run.font.name = MONO_FONT
            run.font.size = Pt(9)
            continue

        # 표 — 헤더 다음 줄이 구분선이면 표로 본다
        if (stripped.startswith("|") and i + 1 < len(lines)
                and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip())):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i].strip()))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for idx, text in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                add_runs(cell.paragraphs[0], text, base_size=9.5)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for idx, text in enumerate(row[:len(header)]):
                    cells[idx].text = ""
                    add_runs(cells[idx].paragraphs[0], text, base_size=9.5)
            doc.add_paragraph()
            continue

        # 수평선 — 문자로 그리면 폭이 안 맞고 글꼴에 따라 끊긴다. 문단 아래 테두리로.
        if re.match(r"^-{3,}$", stripped):
            add_hrule(doc)
            i += 1
            continue

        # 제목
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 4))
            i += 1
            continue

        # 인용 — 연속 줄을 한 문단으로 묶는다
        if stripped.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            add_runs(para, " ".join(q for q in quote if q))
            for run in para.runs:
                run.italic = True
                run.font.color.rgb = GRAY
            continue

        # 목록
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Pt(18 + len(m.group(1)) * 6)
            add_runs(para, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)", line)
        if m:
            para = doc.add_paragraph(style="List Number")
            add_runs(para, m.group(3))
            i += 1
            continue

        # 일반 문단 — 마크다운은 **연속 줄을 한 문단으로 합친다.**
        # 줄마다 문단을 만들면 줄바꿈을 넘어가는 `**굵게**`가 쪼개져 별표가 그대로
        # 남고(실제로 12곳에서 발생), 문단도 어색하게 끊긴다.
        buf = []
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or _is_structural(nxt, lines, i):
                break
            buf.append(nxt)
            i += 1
        para = doc.add_paragraph()
        add_runs(para, " ".join(buf))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(2)
    md = pathlib.Path(sys.argv[1])
    out = pathlib.Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    convert(md, out)
    print("작성: %s (%.1f KB)" % (out, out.stat().st_size / 1024))


if __name__ == "__main__":
    main()
